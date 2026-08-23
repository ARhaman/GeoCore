from pathlib import Path
from docx import Document

source=Path('/home/ubuntu/seg_project/manual_revision_work/MSEGMS_user_upgraded_streamlined_FINAL.docx')
out=Path('/home/ubuntu/seg_project/manual_revision_work/MSEGMS_user_upgraded_streamlined_FINAL_v2.docx')
doc=Document(source)

def append_once(prefix, addition):
    for p in doc.paragraphs:
        if p.text.strip().startswith(prefix) and addition not in p.text:
            p.add_run(addition)
            return True
    return False

# Renumber conclusion after removal of the standalone limitations section.
for p in doc.paragraphs:
    if p.text.strip() == '6. Conclusions':
        p.clear(); p.add_run('5. Conclusions')
        break

# Add explicit first-use calls for the internally aligned supplementary tables.
append_once('Averaged across the ten grouped held-out folds, IoU was', ' Table 4 lists these class-specific grouped-CV mean IoU values.')
append_once('Detailed analysis of two difficult held-out source/lithology cases', ' Table 5 and Figure 6 summarize the two cases; Tables S2 and S3 provide the underlying dominant pixel transitions and semantic-interface distortions.')
append_once('Published mineral-segmentation studies confirm', ' The study-by-study scope comparison is summarized in Table S1.')
# Supplement headings must also be referenced from their related text rather than merely appear in the endmatter.
append_once('The semantic interface analysis showed that composition error also alters labelled boundary structure.', ' The largest source-specific interface distortions are reported in Table S3.')

# Correct any old generic Figure 6 label that survived a user edit.
for p in doc.paragraphs:
    if p.text.strip().startswith('Figure 6. Typical native-resolution'):
        p.clear(); p.add_run('Figure 6. Quantified source-specific B7–B8 error mechanisms. Each row presents a real polarized-light input, reference mask, and ensemble prediction for one held-out source. Source-level mineral mIoU and composition L1 aggregate all eight registered angular views; detailed pixel transitions and semantic-interface distortions are reported in Table 5 and Tables S2–S3.')

doc.save(out)
print(out)
