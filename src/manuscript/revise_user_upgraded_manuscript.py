"""Directly revise the user-upgraded authoritative manuscript without rebuilding from an older draft.

The script preserves the provided DOCX's existing content, tables, figures, author page, and manual
formatting wherever possible. It makes targeted editorial corrections: abbreviations, merged headings,
professional native Word equation blocks, table-number consistency, in-text figure/table references,
limitations within the Discussion, and one consolidated Supplementary Material section.
"""
from __future__ import annotations

import argparse
from copy import deepcopy
from pathlib import Path
from xml.sax.saxutils import escape

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.shared import Pt, RGBColor
from docx.text.paragraph import Paragraph

NS = 'xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"'


def paragraphs(doc):
    return list(doc.paragraphs)


def find_start(doc, prefix: str):
    for p in paragraphs(doc):
        if p.text.strip().startswith(prefix):
            return p
    return None


def set_text(p, text: str, bold: bool = False, italic: bool = False, color=None):
    p.clear()
    run = p.add_run(text)
    run.bold = bold; run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p


def set_heading(p, text: str, level: int = 2):
    p.style = f'Heading {level}'
    set_text(p, text, bold=True, color=(20, 78, 111))
    return p


def insert_after(anchor, text: str = '', style: str | None = None):
    new_p = OxmlElement('w:p')
    anchor._p.addnext(new_p)
    paragraph = Paragraph(new_p, anchor._parent)
    if style:
        paragraph.style = style
    if text:
        paragraph.add_run(text)
    return paragraph


def insert_table_after(doc, anchor, headers, rows, caption: str):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    for cell, label in zip(table.rows[0].cells, headers):
        cell.text = label
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row):
            cell.text = str(value)
    anchor._p.addnext(table._tbl)
    caption_p = insert_after(Paragraph(table._tbl, anchor._parent), caption)
    caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in caption_p.runs:
        run.italic = True; run.font.size = Pt(9)
    return caption_p


def delete_paragraph(p):
    element = p._element
    element.getparent().remove(element)
    p._p = p._element = None


def delete_between(doc, start_prefix: str, end_prefix: str, include_start: bool = False):
    deleting = False
    targets = []
    for p in paragraphs(doc):
        text = p.text.strip()
        if text.startswith(start_prefix):
            deleting = True
            if include_start:
                targets.append(p)
            continue
        if deleting and text.startswith(end_prefix):
            break
        if deleting:
            targets.append(p)
    for p in targets:
        delete_paragraph(p)


def math_para_after(anchor, omml_inner: str, equation_number: int):
    xml = f'<m:oMathPara {NS}><m:oMath>{omml_inner}</m:oMath></m:oMathPara>'
    math = parse_xml(xml)
    anchor._p.addnext(math)
    # Insert equation label as normal paragraph after the math block.
    label_xml = OxmlElement('w:p')
    math.addnext(label_xml)
    label = Paragraph(label_xml, anchor._parent)
    label.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    label.add_run(f'Eq. ({equation_number})').italic = True
    return label


def mrun(text: str) -> str:
    return f'<m:r><m:t>{escape(text)}</m:t></m:r>'


def msub(base: str, sub: str) -> str:
    return f'<m:sSub><m:e>{mrun(base)}</m:e><m:sub>{mrun(sub)}</m:sub></m:sSub>'


def msup(base: str, sup: str) -> str:
    return f'<m:sSup><m:e>{mrun(base)}</m:e><m:sup>{mrun(sup)}</m:sup></m:sSup>'


def mfrac(numerator: str, denominator: str) -> str:
    return f'<m:f><m:num>{mrun(numerator)}</m:num><m:den>{mrun(denominator)}</m:den></m:f>'


def add_professional_equations(doc):
    # Remove current LaTex-looking equation paragraphs in the mathematical section.
    start = find_start(doc, '2.7 Mathematical formulation')
    end = find_start(doc, '2.8 Map-derived')
    if not start or not end:
        return
    replace_prefixes = ('\\mathcal', '\\forall', '\\mathbf', '\\operatorname', '\\mathrm{Dice}', 'For B7, let', 'B8 uses the same', 'The probability ensemble for pixel', 'For mineral class')
    targets = []
    active = False
    for p in paragraphs(doc):
        if p is start:
            active = True
            continue
        if p is end:
            break
        if active and p.text.strip().startswith(replace_prefixes):
            targets.append(p)
    for p in targets:
        delete_paragraph(p)
    # Add a compact native Word-equation set after the heading. The sentences defining variables remain.
    anchor = start
    intro = insert_after(anchor, 'The protocol is summarized by the following equations. All symbols are defined at first use; the equations are native Word mathematical objects.', 'Normal')
    e1 = mrun('𝒟 = { (') + msub('x', 's,v') + mrun(', ') + msub('y', 's') + mrun(') : s ∈ 𝒮, v ∈ ') + msub('𝒱', 's') + mrun(' }')
    e2 = mrun('For every source s, ') + mrun('{') + msub('x', 's,v') + mrun(' : v ∈ ') + msub('𝒱', 's') + mrun('} ⊂ 𝒟train or 𝒟val or 𝒟test')
    e3 = msub('ℒ', 'B7') + mrun(' = ') + msub('ℒ', 'WCE') + mrun(' + (1/8) Σ') + msub('c=1', '8') + mrun(' [1 − Dice') + msub('c', '') + mrun(']')
    e4 = msub('ℒ', 'B8') + mrun(' = 0.7 ') + msub('ℒ', 'WCE') + mrun(' + 0.3 ') + msub('ℒ', 'FT') + mrun('(α=0.35, β=0.65, γ=0.75)')
    e5 = msub('p', 'i') + msup('', 'ens') + mrun(' = w ') + msup(msub('p', 'i'), 'B7') + mrun(' + (1 − w) ') + msup(msub('p', 'i'), 'B8') + mrun(',   ') + msub('ŷ', 'i') + mrun(' = argmax') + msub('c', '') + msup(msub('p', 'i,c'), 'ens')
    e6 = msub('IoU', 'c') + mrun(' = ') + mfrac('TPc', 'TPc + FPc + FNc') + mrun(',   ') + msub('Dice', 'c') + mrun(' = ') + mfrac('2TPc', '2TPc + FPc + FNc')
    e7 = msub('A', 's,c') + mrun(' = ') + mfrac('N(s,c)', 'Ns') + mrun(',   ') + msub('L1', 's') + mrun(' = Σ') + msub('c=1', '8') + mrun(' | ') + msup(msub('A', 's,c'), 'pred') + mrun(' − ') + msup(msub('A', 's,c'), 'GT') + mrun(' |')
    e8 = msub('I', 's,a,b') + mrun(' = (1/8) Σ') + msub('v∈𝒱s', '') + mfrac('E(v,a,b)', 'E(v,min)')
    for number, equation in enumerate((e1,e2,e3,e4,e5,e6,e7,e8), start=1):
        anchor = math_para_after(anchor, equation, number)
    # Replace the old paragraph describing map-derived formulas with a reference to Eqs. (7)–(8).
    map_p = find_start(doc, 'For class c in source s, the two-dimensional')
    if map_p:
        set_text(map_p, 'For class c in source s, the two-dimensional mineral modal-area proxy and composition L1 error are defined in Eq. (7). The normalized mineral-to-mineral semantic pixel-edge interface frequency is defined in Eq. (8). These values are map-derived two-dimensional descriptors, not bulk three-dimensional modal mineralogy or grain-contact topology.')


def update_headings_and_text(doc):
    # Methods merged headings.
    mappings = {
        '2.1 Dataset and semantic-label audit': (2, '2.1 Dataset, registration, and source-level evaluation'),
        '2.2 Source-level partitioning and leakage control': (2, '2.1 Dataset, registration, and source-level evaluation'),
        '2.3 Preprocessing, augmentation, and rare-mineral sampling': (2, '2.2 Preprocessing and B7–B8 model development'),
        '2.4 B7 and B8 segmentation models': (2, '2.2 Preprocessing and B7–B8 model development'),
        '2.5 Validation-calibrated probability ensemble': (2, '2.3 Ensemble, metrics, and statistical inference'),
        '2.6 Metrics and statistical analysis': (2, '2.3 Ensemble, metrics, and statistical inference'),
        '2.7 Mathematical formulation of the source-level segmentation and ensemble protocol': (2, '2.3 Ensemble, metrics, and statistical inference'),
        '2.8 Map-derived quantitative petrography analysis': (2, '2.4 Map-derived quantitative petrography'),
        '3.1 Development progression and selection of the final model': (2, '3.1 Model development, selection, and grouped validation'),
        '3.2 Original source-held test performance': (2, '3.1 Model development, selection, and grouped validation'),
        '3.3 Primary repeated grouped cross-validation estimate': (2, '3.1 Model development, selection, and grouped validation'),
        '3.4 Per-mineral cross-validation behaviour': (2, '3.1 Model development, selection, and grouped validation'),
        '3.5 Native-resolution qualitative evaluation': (2, '3.2 Native-resolution qualitative evaluation'),
        '3.6 Source-held quantitative petrography analysis': (2, '3.3 Map-derived composition and source-specific errors'),
        '3.7 Contrasting source-specific error mechanisms': (2, '3.3 Map-derived composition and source-specific errors'),
        '4.1 What the results establish': (2, '4.1 Key findings and literature context'),
        '4.2 Why the source-level design matters': (2, '4.1 Key findings and literature context'),
        '4.3 Relationship to prior mineral segmentation studies': (2, '4.1 Key findings and literature context'),
        '4.4 Interpretation of model complementarity': (2, '4.1 Key findings and literature context'),
        '4.5 Rare mineral phases and geological interpretability': (2, '4.2 Geological implications and quantitative petrography'),
        '4.6 Implications for quantitative petrography and petroleum-related use': (2, '4.2 Geological implications and quantitative petrography'),
        '4.8 Interpretation of map-derived petrographic quantities': (2, '4.2 Geological implications and quantitative petrography'),
        '4.7 Limitations and next research stage': (2, '4.3 Limitations and next research stage'),
        '6. Conclusions': (1, '5. Conclusions'),
        'Appendix A. Dataset audit and source-level safeguards': (1, 'Supplementary Material'),
        'Appendix B. Figure and reproducibility note': (2, 'S1. Dataset audit and reproducibility'),
        'Appendix C. Quantitative petrography definitions and outputs': (2, 'S2. Quantitative petrography outputs'),
        'Appendix D. Interpretation boundaries': (2, 'S3. Interpretation boundaries'),
    }
    seen = set()
    for p in paragraphs(doc):
        text = p.text.strip()
        for old, (level, new) in mappings.items():
            if text == old:
                # Multiple old headings can map to one new heading; retain only the first visible instance.
                if new in seen:
                    delete_paragraph(p)
                else:
                    set_heading(p, new, level); seen.add(new)
                break
    # Remove only the standalone limitations section; limitations already reside at the end of Discussion.
    # The conclusion heading has been renumbered above from 6 to 5, so it is the stop marker.
    delete_between(doc, '5. Limitations', '5. Conclusions', include_start=True)

    # Correct duplicate table numbering and clarify supplement location.
    replacements = {
        'Table 4. Source-held two-dimensional mineral modal-area proxies': 'Table 6. Source-held two-dimensional mineral modal-area proxies',
        'Full matrices are provided in Appendix D': 'Full matrices are provided in Supplementary Material S3',
        'Figures 5–7 should present real, native-resolution': 'Figure 5 presents real, native-resolution',
        'The strong cases show recovery': 'Figure 5 shows recovery',
        'Strong cases show coherent': 'Figure 5 shows coherent',
        'The qualitative figures are therefore essential': 'The qualitative figures are therefore essential',
    }
    for p in paragraphs(doc):
        for old, new in replacements.items():
            if old in p.text:
                set_text(p, p.text.replace(old, new))


def add_in_text_references(doc):
    # Add precise first-use cross-references without altering the existing captions.
    targets = [
        ('We therefore define one thin section as one independent geological source.', ' The dataset construction and representative real image–mask diversity are shown in Figures 1 and 2, respectively.'),
        ('The final model was a probability-level ensemble of B7 and B8.', ' Its architecture, probability fusion, and validation-only selection are summarized in Figure 3.'),
        ('The early compact U-Net baseline (B0) achieved', ' The complete original fixed-split development progression is reported in Table 1 and Figure 4A.'),
        ('On the 10-source original held-out test partition,', ' The class-level original held-out results are listed in Table 2.'),
        ('The primary estimate of model generalization was obtained', ' The grouped source-level summary and per-mineral behaviour are reported in Tables 3–4 and Figure 4B–C.'),
        ('Figure 5 presents real, native-resolution', ' The selected illustrative case is shown in Figure 5; it is not the primary accuracy estimate.'),
        ('Detailed analysis of two difficult held-out source/lithology cases', ' Table 5 and Figure 6 present the aligned source-specific comparison; Tables S2–S3 provide the detailed confusion and interface pathways.'),
        ('The modal-area and interface analyses processed', ' Table 6 reports mineral-specific 2D composition bias; the full semantic-interface matrix is retained in Supplementary Material S3.'),
    ]
    for p in paragraphs(doc):
        for prefix, suffix in targets:
            if p.text.startswith(prefix) and suffix.strip() not in p.text:
                set_text(p, p.text + suffix)


def insert_abbreviation_table(doc):
    keyword = find_start(doc, 'Keywords:')
    if not keyword:
        return
    # Do not duplicate on repeat runs.
    if any(p.text.strip() == 'Abbreviations' for p in paragraphs(doc)):
        return
    heading = insert_after(keyword, 'Abbreviations', 'Heading 2')
    rows = [
        ['ASPP', 'Atrous spatial pyramid pooling'],
        ['B7 / B8', 'Final ResNet-34 U-Net / DeepLabV3–ResNet50 model variants'],
        ['CV', 'Cross-validation'],
        ['Dice', 'Dice similarity coefficient'],
        ['FN / FP / TP', 'False negative / false positive / true positive pixels'],
        ['IoU / mIoU', 'Intersection-over-union / mean mineral intersection-over-union'],
        ['L1', 'Sum of absolute mineral-area proxy errors'],
        ['RGB', 'Red–green–blue image channels'],
        ['SIFT', 'Scale-invariant feature transform'],
        ['WCE', 'Class-weighted cross-entropy'],
    ]
    insert_table_after(doc, heading, ['Abbreviation', 'Definition'], rows, 'Abbreviations used in the manuscript. Symbols used only in equations are defined at first use.')


def normalize_reference_labels(doc):
    # Resolve duplicate/non-sequential labels introduced during manual upgrades while preserving reference text.
    mapping = [
        ('[1] Zhu J, Yang J.', '[1] Zhu J, Yang J.'),
        ('[2] Saxena', '[2] Saxena'),
        ('[3] Hassan', '[3] Hassan'),
        ('[4] Acevedo', '[4] Acevedo'),
        ('[5] Chen Z', '[5] Chen Z'),
        ('[6] Pirrie', '[6] Pirrie'),
        ('[8] Tang', '[7] Tang'),
        ('Lawrence M. Anovitz', '[8] Lawrence M. Anovitz'),
        ('[9] Zhong', '[9] Zhong'),
        ('[7] He K', '[10] He K'),
        ('[8] Chen L-C', '[11] Chen L-C'),
    ]
    for p in paragraphs(doc):
        text = p.text.strip()
        for start, repl in mapping:
            if text.startswith(start):
                if start == 'Lawrence M. Anovitz':
                    set_text(p, repl + text)
                else:
                    set_text(p, repl + text[len(start):])
                break


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument('--input', required=True)
    parser.add_argument('--output', required=True)
    args = parser.parse_args()
    doc = Document(args.input)
    insert_abbreviation_table(doc)
    update_headings_and_text(doc)
    add_professional_equations(doc)
    add_in_text_references(doc)
    normalize_reference_labels(doc)
    output = Path(args.output); output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    print(output)

if __name__ == '__main__':
    main()
