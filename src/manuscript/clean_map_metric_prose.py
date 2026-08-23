from pathlib import Path
from docx import Document

source=Path('/home/ubuntu/seg_project/manual_revision_work/MSEGMS_user_upgraded_revised_v4.docx')
out=Path('/home/ubuntu/seg_project/manual_revision_work/MSEGMS_user_upgraded_revised_v5.docx')
doc=Document(source)
replacements={
    'For class c in source s, the two-dimensional mineral modal-area proxy was': 'For mineral class c in source s, the two-dimensional modal-area proxy Aₛ,꜀ and source-level composition L1 discrepancy are defined in Eq. (7). Each term is computed from predicted or reference pixels within a registered view, averaged across the eight views of a source, and summarized across the 10 held-out sources. The resulting quantities are two-dimensional map-derived area proxies, not estimates of bulk three-dimensional modal mineralogy.',
    'We also counted unlike four-neighbour pixel edges between mineral labels a and b.': 'The normalized four-neighbour semantic pixel-edge interface frequency Iₛ,ₐ,ᵦ is defined in Eq. (8). It is calculated only for unlike mineral labels, normalized by all mineral-to-mineral unlike edges within each view, and then aggregated across the eight registered views of a source. This diagnostic does not delineate individual grains, quantify physical three-dimensional contact area, or establish a grain-contact network [2,4].',
}
for p in doc.paragraphs:
    for prefix,new in replacements.items():
        if p.text.strip().startswith(prefix):
            p.clear(); p.add_run(new)
            break
doc.save(out)
print(out)
