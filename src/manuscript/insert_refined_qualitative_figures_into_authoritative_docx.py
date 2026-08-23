from __future__ import annotations
import shutil, zipfile
from pathlib import Path
from tempfile import TemporaryDirectory
from docx import Document

SOURCE=Path('/home/ubuntu/seg_project/manual_revision_work/MSEGMS_user_upgraded_revised_v5.docx')
OUT=Path('/home/ubuntu/seg_project/manual_revision_work/MSEGMS_user_upgraded_streamlined_FINAL.docx')
FIG5=Path('/home/ubuntu/seg_project/manual_revision_work/refined_word_figures/Figure_5_strong_native_ensemble_900dpi.png')
FIG6=Path('/home/ubuntu/seg_project/manual_revision_work/refined_word_figures/Figure_6_source_specific_errors_900dpi.png')

with TemporaryDirectory() as temp:
    root=Path(temp)/'docx'; root.mkdir()
    with zipfile.ZipFile(SOURCE) as archive: archive.extractall(root)
    media=root/'word/media'
    shutil.copy2(FIG5,media/'image5.png')
    shutil.copy2(FIG6,media/'image6.png')
    packed=Path(temp)/'manuscript.docx'
    with zipfile.ZipFile(packed,'w',zipfile.ZIP_DEFLATED) as archive:
        for item in root.rglob('*'):
            if item.is_file(): archive.write(item,item.relative_to(root))
    shutil.copy2(packed,OUT)

doc=Document(OUT)
replacements={
    'Figure 5. Strong native-resolution B7–B8 ensemble predictions.': 'Figure 5. Strong held-out B7–B8 ensemble illustration. The row presents a real polarized-light input, its reference semantic mask, and the validation-calibrated B7–B8 prediction for Muscovite–Granite–3. Frame mineral mIoU is displayed only to identify the illustration; the primary accuracy evidence remains the repeated grouped source-level cross-validation estimate.',
    'Figure 6. Typical native-resolution B7–B8 ensemble predictions.': 'Figure 6. Quantified source-specific B7–B8 error mechanisms. Each row presents a real polarized-light input, reference mask, and ensemble prediction for one held-out source. Source-level mineral mIoU and composition L1 aggregate all eight registered angular views; detailed pixel transitions and semantic-interface distortions are reported in Table 5 and Tables S2–S3.',
    'Figure 5 presents real, native-resolution 1280 × 1024 microscopy frames, their semantic ground-truth masks, and B7–B8 ensemble predictions for strong, typical, and difficult cases, respectively.': 'Figures 5–7 present real native-resolution microscopy inputs, semantic ground-truth masks, and B7–B8 ensemble predictions. Figure 5 provides one clearly labelled favourable held-out illustration; Figure 6 links two difficult source cases directly to the measured source-level errors in Table 5; and Figure 7 retains an additional difficult qualitative example to avoid selective reporting.',
    'Figure 5 shows recovery of broad mineral domains and coherent grain-scale structures. Typical cases show successful localization of large regions but imperfect mineral identity assignment at adjacent phases and complex boundaries. Difficult cases, including the Limburgite examples, demonstrate missed small regions, overprediction, and confusion among visually similar mineral phases. Showing all three conditions prevents selective visualization from overstating model performance. The primary quantitative claim remains the repeated grouped source-level cross-validation estimate in Table 3.': 'Figure 5 shows recovery of broad mineral domains in a favourable held-out case. Figure 6 makes the complementary limitation explicit: Muscovite–Schist–3 shows severe quartz–muscovite label fragmentation, whereas Pyroxene–Peridotite–1 shows dominant pyroxene-to-olivine reassignment. Figure 7 retains an additional difficult case. Together, these panels prevent selective visualization from overstating the model; the primary quantitative claim remains the repeated grouped source-level cross-validation estimate in Table 3.',
}
for p in doc.paragraphs:
    for old,new in replacements.items():
        if p.text.strip().startswith(old):
            p.clear(); p.add_run(new); break
# Preserve Figure 7 but place it in the supplementary interpretation by a clarifying caption.
for p in doc.paragraphs:
    if p.text.strip().startswith('Figure 7. Difficult native-resolution'):
        p.clear(); p.add_run('Figure 7. Additional difficult native-resolution B7–B8 ensemble example. This supplementary-style main-text panel shows a real input, reference mask, and prediction for a low-performing source; it is retained to document failure modes beyond the quantified cases in Figure 6.')
        break
doc.save(OUT)
print(OUT)
