from __future__ import annotations

import argparse
import shutil
import zipfile
from pathlib import Path
from tempfile import TemporaryDirectory

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from docx.shared import Pt


def find(doc, prefix):
    for p in doc.paragraphs:
        if p.text.strip().startswith(prefix):
            return p
    return None


def delete_paragraph(p):
    element = p._element
    element.getparent().remove(element)


def insert_after(anchor, text='', style=None):
    new = OxmlElement('w:p'); anchor._p.addnext(new)
    p = Paragraph(new, anchor._parent)
    if style: p.style = style
    if text: p.add_run(text)
    return p


def replace_ensemble_section(doc):
    start = find(doc, '2.3 Ensemble, metrics, and statistical inference')
    end = find(doc, '2.4 Map-derived quantitative petrography')
    if not start or not end:
        raise RuntimeError('Expected streamlined Section 2.3/2.4 headings were not found.')
    between=[]; active=False
    for p in list(doc.paragraphs):
        if p is start: active=True; continue
        if p is end: break
        if active: between.append(p)
    for p in between: delete_paragraph(p)
    anchor=start
    paragraphs=[
        'B7 was a pretrained ResNet-34 U-Net trained with class-weighted cross-entropy plus mineral Dice loss. B8 was a pretrained DeepLabV3–ResNet50 model trained with class-weighted cross-entropy plus focal Tversky loss. Both models used the same normalized labels, source-level split, crop geometry, rare-mineral crop policy, and held-out evaluation protocol. Figure 3 summarizes their complementary branches and validation-only probability fusion.',
        'Let 𝒮 denote the set of thin-section sources and 𝒱ₛ the eight registered views of source s. A source, rather than an individual angular frame, is the independent evaluation unit. Equations (1)–(6) define the source-level split, ensemble rule, and segmentation metrics; Eqs. (7)–(8) define the supplementary map-derived quantities.',
    ]
    for text in paragraphs:
        anchor=insert_after(anchor,text,'Normal')
    equations=[
        '𝒟 = {(xₛ,ᵥ, yₛ) : s ∈ 𝒮, v ∈ 𝒱ₛ}',
        '∀ s ∈ 𝒮 : {xₛ,ᵥ : v ∈ 𝒱ₛ} ⊂ 𝒟train  or  𝒟val  or  𝒟test',
        'ℒB7 = ℒWCE + (1/8) ∑₍c₌₁₎⁸ [1 − Dicec]',
        'ℒB8 = 0.7 ℒWCE + 0.3 ℒFT(α = 0.35, β = 0.65, γ = 0.75)',
        'pᵉⁿˢᵢ = w pᴮ⁷ᵢ + (1 − w) pᴮ⁸ᵢ,    ŷᵢ = arg max₍c₎ pᵉⁿˢᵢ,꜀',
        'IoUc = TPc / (TPc + FPc + FNc),    Dicec = 2TPc / (2TPc + FPc + FNc)',
        'Aₛ,꜀ = Nₛ,꜀ / Nₛ,    L1ₛ = ∑₍c₌₁₎⁸ |Aᵖʳᵉᵈₛ,꜀ − Aᴳᵀₛ,꜀|',
        'Iₛ,ₐ,ᵦ = (1/8) ∑₍v∈𝒱ₛ₎ E(v,a,b) / E(v,mineral)',
    ]
    for number, equation in enumerate(equations, 1):
        p=insert_after(anchor,'','Normal'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        r=p.add_run(equation + f'     ({number})'); r.font.name='Cambria Math'; r.font.size=Pt(11)
        anchor=p
    anchor=insert_after(anchor,'In Eq. (7), A is a two-dimensional mineral modal-area proxy and L1 is the source-level composition discrepancy. In Eq. (8), I is a normalized four-neighbour semantic pixel-edge interface frequency. These quantities are not three-dimensional modal mineralogy, grain-contact topology, porosity, or physical-scale measurements.','Normal')
    return doc


def font(size, bold=False):
    choices = [
        '/usr/share/fonts/truetype/msttcorefonts/Arial.ttf',
        '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf',
    ]
    if bold:
        choices = ['/usr/share/fonts/truetype/msttcorefonts/Arial_Bold.ttf','/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf'] + choices
    for path in choices:
        if Path(path).exists(): return ImageFont.truetype(path,size)
    return ImageFont.load_default()


def box(draw, xy, header, body, fill, border, head_fill):
    x1,y1,x2,y2=xy
    draw.rounded_rectangle(xy, radius=24, fill=fill, outline=border, width=5)
    draw.rounded_rectangle((x1,y1,x2,y1+64), radius=24, fill=head_fill)
    draw.rectangle((x1,y1+40,x2,y1+64), fill=head_fill)
    hfont=font(27,True); bfont=font(32,True)
    draw.text(((x1+x2)//2,y1+20),header,font=hfont,fill='white',anchor='mm')
    lines=body.split('\n'); start=(y1+y2)//2-(len(lines)-1)*22
    for i,line in enumerate(lines): draw.text(((x1+x2)//2,start+i*46),line,font=bfont,fill='#122534',anchor='mm')


def arrow(draw, start, end, color='#147a70'):
    draw.line([start,end],fill=color,width=10)
    import math
    angle=math.atan2(end[1]-start[1],end[0]-start[0]); length=28
    a1=(end[0]-length*math.cos(angle-0.55),end[1]-length*math.sin(angle-0.55)); a2=(end[0]-length*math.cos(angle+0.55),end[1]-length*math.sin(angle+0.55))
    draw.polygon([end,a1,a2],fill=color)


def generate_figure(path):
    W,H=6000,3375; im=Image.new('RGB',(W,H),'white'); d=ImageDraw.Draw(im)
    navy='#14506f'; teal='#168272'; blue='#3d82b3'; rust='#b9644e'; light='#eef5f7'; dark='#142633'
    d.rectangle((0,0,W,260),fill=navy)
    d.text((150,88),'Figure 3. B7–B8 probability ensemble and validation-only selection',font=font(76,True),fill='white')
    d.text((150,305),'Single-view mineral semantic segmentation with complementary losses, probability fusion, and strictly separated model selection.',font=font(35),fill='#4d5a63')
    input_box=(180,730,1200,1730); b7=(1600,610,3100,1170); b8=(1600,1430,3100,1990); ens=(3500,860,4760,1740); val=(5100,610,5820,1170); cv=(5100,1430,5820,1990)
    box(d,input_box,'REGISTERED INPUT','One angular RGB view\n640 × 512 normalized',light,teal,teal)
    box(d,b7,'B7  PRETRAINED RESNET-34 U-NET','Weighted cross-entropy\n+ mineral Dice loss',light,blue,blue)
    box(d,b8,'B8  DEEPLABV3–RESNET50','Weighted cross-entropy\n+ focal Tversky loss',light,rust,rust)
    box(d,ens,'PROBABILITY ENSEMBLE','pᵉⁿˢᵢ = w pᴮ⁷ᵢ + (1 − w) pᴮ⁸ᵢ\nŷᵢ = arg max₍c₎ pᵉⁿˢᵢ,꜀', '#e8f4f3',teal,'#e8f4f3')
    box(d,val,'VALIDATION-ONLY CALIBRATION','Original fixed split\nB7 = 0.6; B8 = 0.4',light,navy,navy)
    box(d,cv,'GROUPED-CV RESELECTION','Fold B7 weights: 0.3–0.6\nmean = 0.44',light,teal,teal)
    arrow(d,(1200,1120),(1600,890)); arrow(d,(1200,1330),(1600,1710)); arrow(d,(3100,890),(3500,1120)); arrow(d,(3100,1710),(3500,1480)); arrow(d,(4760,1120),(5100,890)); arrow(d,(4760,1480),(5100,1710))
    d.rounded_rectangle((190,2410,5810,2780),radius=25,fill='#f5f8fa',outline=rust,width=5)
    d.text((3000,2590),'Model-selection safeguard: test sources never selected a checkpoint, crop policy, loss configuration, or ensemble coefficient.',font=font(39,True),fill=dark,anchor='mm')
    d.text((3000,3070),'Equations (1)–(6) define the source-level split, loss terms, ensemble, and segmentation metrics in Section 2.3.',font=font(28),fill='#56666f',anchor='mm')
    im.save(path,dpi=(900,900))


def replace_media(docx_path, media_png):
    with TemporaryDirectory() as temp:
        root=Path(temp)/'docx'; root.mkdir()
        with zipfile.ZipFile(docx_path) as z: z.extractall(root)
        target=root/'word/media/image3.png'
        if not target.exists(): raise RuntimeError('Expected Figure 3 media image3.png was not found.')
        shutil.copy2(media_png,target)
        temp_docx=Path(temp)/'repacked.docx'
        with zipfile.ZipFile(temp_docx,'w',zipfile.ZIP_DEFLATED) as z:
            for file in root.rglob('*'):
                if file.is_file(): z.write(file,file.relative_to(root))
        shutil.copy2(temp_docx,docx_path)


def main():
    p=argparse.ArgumentParser(); p.add_argument('--input',required=True); p.add_argument('--output',required=True); p.add_argument('--figure-output',required=True)
    a=p.parse_args(); source=Path(a.input); output=Path(a.output); output.parent.mkdir(parents=True,exist_ok=True)
    doc=Document(source); replace_ensemble_section(doc); doc.save(output)
    fig=Path(a.figure_output); fig.parent.mkdir(parents=True,exist_ok=True); generate_figure(fig); replace_media(output,fig)
    print(output); print(fig)

if __name__=='__main__': main()
